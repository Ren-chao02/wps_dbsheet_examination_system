#!/usr/bin/env python3
"""Generate wps_table.docx from the design plan."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# -- Page setup --
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# -- Style tweaks --
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for level in range(1, 4):
    heading = doc.styles[f'Heading {level}']
    heading.font.name = '微软雅黑'
    heading.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    heading.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

# -- Cover / Title --
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('金山多维表格考试系统\n实施方案')
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(0x66, 0x7E, 0xEA)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('—— 设计文档 v1.0 ——')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('2024年6月').font.size = Pt(11)

doc.add_page_break()

# -- TOC placeholder --
doc.add_heading('目录', level=1)
doc.add_paragraph('（在 Word 中按 Ctrl+A → F9 刷新目录域）')
doc.add_paragraph()
# Insert TOC field
paragraph = doc.add_paragraph()
run = paragraph.add_run()
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
run._r.append(fldChar1)

instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = ' TOC \\o "1-3" \\h \\z '
run._r.append(instrText)

fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')
run._r.append(fldChar2)

run2 = paragraph.add_run('（右键点击此处 → 更新域 → 更新整个目录）')
run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

fldChar3 = OxmlElement('w:fldChar')
fldChar3.set(qn('w:fldCharType'), 'end')
run._r.append(fldChar3)

doc.add_page_break()

# ============================================================
# HELPER: add code block as formatted paragraph
# ============================================================
def add_code_block(text, lang=''):
    """Add a monospaced code block."""
    for line in text.strip().split('\n'):
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.2
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        # Light gray background
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F5F5F5')
        shading.set(qn('w:val'), 'clear')
        p.paragraph_format.element.get_or_add_pPr().append(shading)

def add_table(headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()

# ============================================================
# CONTENT SECTIONS
# ============================================================

# --- Context ---
doc.add_heading('1. 项目背景 (Context)', level=1)
doc.add_paragraph(
    '构建一个面向金山多维表格操作技能的考核系统。学生在真实金山多维表格环境中完成操作任务'
    '（建表、字段配置、视图管理、表单创建等），系统通过 API + 规则引擎自动验证操作结果并判分。'
    '目标用户为教师（出题、阅卷）、学生（考试、查看成绩）和管理员（用户与系统管理）。'
)

# --- Architecture ---
doc.add_heading('2. 总体架构', level=1)
doc.add_paragraph('系统采用经典的三层架构：前端 React SPA → 后端 Node.js/Express REST API → PostgreSQL + Redis 数据层，外部对接金山多维表格 Open API。')

add_code_block('''
┌─────────────────────────────────────────────────┐
│               前端 (React SPA)                    │
│  ┌──────────┐ ┌──────────┐ ┌──────────────┐     │
│  │ 教师工作台│ │ 学生考试端│ │ 管理后台     │     │
│  └──────────┘ └──────────┘ └──────────────┘     │
└─────────────────────┬───────────────────────────┘
                      │ HTTP REST + WebSocket
┌─────────────────────▼───────────────────────────┐
│            后端服务 (Node.js + Express)           │
│  ┌────────┐ ┌────────┐ ┌────────┐ ┌─────────┐  │
│  │用户认证 │ │题库管理 │ │考试引擎 │ │判分引擎  │  │
│  └────────┘ └────────┘ └────────┘ └─────────┘  │
│  ┌──────────────────────────────────────────┐   │
│  │   验证适配层 (可插拔)                      │   │
│  │   API验证器 | 快照验证器 | 人工复核         │   │
│  └──────────────────────────────────────────┘   │
└─────────────────────┬───────────────────────────┘
                      │
         ┌────────────┴────────────┐
         │  PostgreSQL  |  Redis   │
         └─────────────────────────┘
                      │
         ┌────────────┴────────────┐
         │  金山多维表格 Open API   │
         └─────────────────────────┘
''')

# --- Tech Stack ---
doc.add_heading('3. 技术栈', level=1)
add_table(
    ['层次', '技术选型', '理由'],
    [
        ['前端框架', 'React 18 + TypeScript + Ant Design 5', 'Table/Form 组件丰富，适合管理后台'],
        ['状态管理', 'Zustand', '轻量，TypeScript 友好'],
        ['路由', 'React Router v6', '标准方案'],
        ['后端', 'Node.js + Express + TypeScript', '前后端统一语言，判分引擎灵活'],
        ['ORM', 'Prisma', '类型安全，迁移方便'],
        ['数据库', 'PostgreSQL 15', 'JSONB 存储题目规则，事务可靠'],
        ['缓存/队列', 'Redis', '考试状态缓存 + 判分任务队列'],
        ['实时通信', 'Socket.IO', '考试监控、倒计时同步'],
        ['容器化', 'Docker Compose', '一键部署'],
    ]
)

# --- Database Schema ---
doc.add_heading('4. 数据库设计', level=1)

doc.add_heading('4.1 核心实体关系', level=2)
doc.add_paragraph('系统核心实体及关联关系如下：')
add_code_block('''
User ──< ExamSession >── Exam
  │                        │
  └──< Question >──────────┘
         │
         └──< VerificationRule (JSONB)

Exam ──< ExamQuestion >── Question
  │
  └──< StudentSubmission >──< SubmissionDetail >
                                  │
                                  └──< VerificationResult >
''')

doc.add_heading('4.2 表结构详述', level=2)

tables_desc = [
    ('users (用户表)', 'id, username, password_hash, real_name, role(admin|teacher|student), email, avatar_url, created_at, updated_at'),
    ('question_categories (题库分类)', 'id, name, parent_id(自引用树形), sort_order, created_at'),
    ('questions (题目)', 'id, category_id, title, description, type(create_table|add_field|config_view|create_form|comprehensive), difficulty, score, answer_rules(JSONB), hints, tags[], status, created_by, created_at, updated_at'),
    ('exams (考试)', 'id, title, description, mode(practice|quiz|exam), duration_minutes, start_time, end_time, total_score, pass_score, status(draft→published→in_progress→ended), settings(JSONB), created_by'),
    ('exam_questions (考试-题目关联)', 'id, exam_id, question_id, sort_order, score_override. UNIQUE(exam_id, question_id)'),
    ('student_submissions (学生答卷)', 'id, exam_id, student_id, table_space_id, status(pending→in_progress→submitted→grading→graded), started_at, submitted_at, graded_at, total_score, grader_comment, graded_by. UNIQUE(exam_id, student_id)'),
    ('submission_details (答题详情)', 'id, submission_id, question_id, answer_json(JSONB), score, is_correct. UNIQUE(submission_id, question_id)'),
    ('verification_results (验证结果)', 'id, submission_detail_id, submission_id, rule_id, action, expected, actual, passed, score, error_message, needs_review, verified_at'),
    ('exam_sessions (考试会话)', 'id, submission_id(UNIQUE), student_id, exam_id, ws_connected, last_heartbeat, ip_address'),
]
for tname, tfields in tables_desc:
    p = doc.add_paragraph()
    run = p.add_run(f'{tname}  ')
    run.font.bold = True
    p.add_run(tfields)

doc.add_heading('4.3 验证规则 JSONB 结构示例', level=2)
add_code_block('''[
  { "id":"r1", "action":"check_table_exists", "params":{"tableName":"学生档案"}, "score":5 },
  { "id":"r2", "action":"check_field",
    "params":{"tableName":"学生档案","fieldName":"姓名","type":"text","required":true},
    "score":5 },
  { "id":"r3", "action":"check_field",
    "params":{"tableName":"学生档案","fieldName":"性别","type":"single_select","options":["男","女"]},
    "score":5 },
  { "id":"r4", "action":"check_view_exists",
    "params":{"tableName":"学生档案","viewType":"kanban","groupByField":"性别"},
    "score":10 }
]''', 'json')

# --- API Design ---
doc.add_heading('5. API 设计', level=1)

api_modules = [
    ('5.1 认证 (Auth)', [
        ('POST /api/auth/login', '登录，返回 JWT'),
        ('POST /api/auth/register', '注册'),
        ('POST /api/auth/refresh', '刷新令牌'),
        ('GET /api/auth/me', '获取当前用户信息'),
    ]),
    ('5.2 用户管理 (Admin)', [
        ('GET /api/users', '用户列表（分页、搜索）'),
        ('POST /api/users', '创建用户'),
        ('PUT /api/users/:id', '更新用户'),
        ('DELETE /api/users/:id', '删除用户'),
        ('POST /api/users/batch-import', '批量导入（CSV）'),
    ]),
    ('5.3 题库管理 (Teacher)', [
        ('GET /api/questions', '题目列表（筛选：分类、难度、类型、状态）'),
        ('POST /api/questions', '创建题目（含验证规则）'),
        ('GET /api/questions/:id', '题目详情'),
        ('PUT /api/questions/:id', '更新题目'),
        ('DELETE /api/questions/:id', '删除题目'),
        ('GET /api/categories', '分类树'),
        ('POST /api/categories', '创建分类'),
    ]),
    ('5.4 考试管理 (Teacher)', [
        ('GET /api/exams', '考试列表'),
        ('POST /api/exams', '创建考试'),
        ('PUT /api/exams/:id/questions', '选题组卷'),
        ('POST /api/exams/:id/publish', '发布考试'),
        ('POST /api/exams/:id/start', '开始考试'),
        ('POST /api/exams/:id/end', '结束考试'),
        ('GET /api/exams/:id/submissions', '查看所有答卷'),
    ]),
    ('5.5 考试参与 (Student)', [
        ('GET /api/my-exams', '我的考试列表'),
        ('GET /api/my-exams/:id', '考试详情/题目'),
        ('POST /api/my-exams/:id/start', '开始答题'),
        ('POST /api/my-exams/:id/submit', '提交答卷'),
        ('GET /api/my-exams/:id/result', '查看成绩'),
    ]),
    ('5.6 判分 (Teacher + Auto)', [
        ('POST /api/grading/:submissionId', '触发自动判分'),
        ('POST /api/grading/:submissionId/detail/:detailId', '手动对某题打分'),
        ('POST /api/grading/:submissionId/finalize', '完成评分'),
        ('POST /api/grading/batch/:examId', '批量评分'),
    ]),
    ('5.7 统计分析', [
        ('GET /api/statistics/overview', '总览仪表盘'),
        ('GET /api/statistics/exam/:examId', '单场考试分析'),
        ('GET /api/statistics/student/:studentId', '学生能力画像'),
    ]),
    ('5.8 金山API代理', [
        ('POST /api/kingsoft/proxy', '代理调用，统一管理Token'),
        ('GET /api/kingsoft/table/:space_id/tables', '获取表列表'),
        ('GET /api/kingsoft/table/:space_id/:table_id/fields', '获取字段'),
        ('GET /api/kingsoft/table/:space_id/:table_id/views', '获取视图'),
    ]),
]
for module_title, routes in api_modules:
    doc.add_heading(module_title, level=2)
    add_table(['方法/路径', '说明'], [(m, d) for m, d in routes])

# --- Verification Engine ---
doc.add_heading('6. 验证引擎设计', level=1)

doc.add_heading('6.1 架构', level=2)
doc.add_paragraph(
    '验证引擎由 GradingService 统一调度，RuleInterpreter 解析题目的 answer_rules JSONB，'
    'VerifierDispatcher 根据 rule.action 路由到对应验证器（API验证器、快照验证器、人工复核），'
    '各验证器实现统一的 IVerifier 接口。'
)

doc.add_heading('6.2 验证动作枚举', level=2)
add_table(
    ['分类', '动作', '说明'],
    [
        ['表操作', 'check_table_exists', '表是否存在'],
        ['表操作', 'check_table_name', '表名称（模糊匹配）'],
        ['表操作', 'check_table_count', '表数量'],
        ['字段', 'check_field', '字段存在、类型、选项'],
        ['字段', 'check_field_count', '字段数量'],
        ['字段', 'check_field_required', '必填设置'],
        ['字段', 'check_field_formula', '公式字段验证'],
        ['字段', 'check_linked_record', '关联记录'],
        ['视图', 'check_view_exists', '视图是否存在'],
        ['视图', 'check_view_type', '视图类型(gird/kanban/gallery/form)'],
        ['视图', 'check_view_filter', '筛选条件'],
        ['视图', 'check_view_sort', '排序规则'],
        ['视图', 'check_view_group', '分组字段'],
        ['表单', 'check_form_exists', '表单是否存在'],
        ['表单', 'check_form_fields', '表单字段可见性'],
        ['表单', 'check_form_settings', '表单设置'],
        ['记录', 'check_record_exists', '记录存在'],
        ['记录', 'check_record_value', '记录值'],
        ['记录', 'check_record_count', '记录数量'],
    ]
)

doc.add_heading('6.3 验证器接口', level=2)
add_code_block('''interface IVerifier {
  name: string;
  verify(rule: VerificationRule, context: VerifierContext): Promise<VerificationResult>;
}

interface VerificationResult {
  ruleId: string;
  passed: boolean;
  score: number;
  expected: any;
  actual: any;
  errorMessage?: string;
  suggestion?: string;
}''', 'typescript')

doc.add_heading('6.4 验证策略', level=2)
strategies = [
    '存在性检查（table/field/view 是否存在）→ API 精确匹配',
    '属性检查（字段类型、视图配置）→ API 读取 + JSON 深度对比',
    '模糊匹配（公式字段内容）→ 正则表达式 / Levenshtein 编辑距离',
    '人工兜底 → 无法自动判断的规则标记 needs_review=true，推送教师复核',
]
for s in strategies:
    doc.add_paragraph(s, style='List Bullet')

# --- Frontend Routes ---
doc.add_heading('7. 前端路由设计', level=1)
add_code_block('''# 认证
/login                          # 登录
/register                       # 注册

# 学生端
/student/dashboard              # 我的考试列表
/student/exam/:id               # 考试说明页
/student/exam/:id/doing         # 答题页（题目导航 + 金山表格跳转）
/student/exam/:id/result        # 成绩页

# 教师端
/teacher/dashboard              # 工作台（统计概览）
/teacher/questions              # 题库管理
/teacher/questions/new          # 新建题目
/teacher/questions/:id/edit     # 编辑题目（含规则构建器）
/teacher/exams                  # 考试管理列表
/teacher/exams/new              # 创建考试
/teacher/exams/:id/edit         # 编辑考试 / 选题组卷
/teacher/exams/:id/monitor      # 考场监控（实时状态）
/teacher/exams/:id/grading      # 阅卷页
/teacher/exams/:id/statistics   # 成绩分析

# 管理员端
/admin/users                    # 用户管理''')

# --- Component Tree ---
doc.add_heading('8. 前端组件树', level=1)
p = doc.add_paragraph('核心组件层级结构：')
add_code_block('''App
├── AuthLayout (LoginPage, RegisterPage)
├── StudentLayout
│   ├── StudentDashboard (ExamCard[])
│   ├── ExamIntroPage
│   ├── ExamDoingPage (ExamTimer, QuestionNav, QuestionCard, SubmitConfirm)
│   └── ExamResultPage (ScoreOverview, QuestionResultList, QuestionResultItem)
├── TeacherLayout
│   ├── TeacherDashboard (StatCard[], RecentExamList)
│   ├── QuestionBank (QuestionFilterBar, QuestionTable)
│   ├── QuestionEditor (BasicInfoForm, RuleBuilder★, QuestionPreview)
│   │   └── RuleBuilder (RuleItem, ActionSelector, ParamEditor, ScoreInput)
│   ├── ExamForm (BasicInfo, QuestionPicker, ExamSettings)
│   ├── ExamMonitor (StudentStatusList, RealTimeAlerts)
│   ├── GradingPage (SubmissionList, GradingDetail, ScoreOverride)
│   └── StatisticsPage (ScoreDistribution, QuestionAnalysis, ExportReport)
├── AdminLayout (UserManagement → UserTable, BatchImportModal)
└── Shared (RichTextEditor, StatusBadge, ConfirmModal, FileUpload)''')

# --- Exam Lifecycle ---
doc.add_heading('9. 考试生命周期', level=1)
doc.add_paragraph('一次完整考试经历以下阶段：')
add_code_block('''
[教师]                  [系统]                    [学生]
  │                       │                         │
  ├─ 创建题目/规则 ──────►│                         │
  ├─ 创建考试/选题 ──────►│                         │
  ├─ 发布考试 ───────────►├─ 通知学生 ────────────►├─ 查看说明
  │                       │◄── 开始答题 ───────────┤
  │                       ├─ 创建答题记录           ├─ 金山表格操作
  │                       ├─ 分配表格空间           ├─ 逐题完成
  │◄─ 实时监控 ──────────┤◄── 心跳上报 ──────────┤
  │                       │◄── 提交答卷 ───────────┤
  │                       ├─ 加入判分队列           │
  │                       ├─ 执行自动判分           │
  │◄─ 人工复核 ──────────┤                         ├─ 查看成绩
  ├─ 修正分数/评语 ──────►├─ 更新最终成绩 ────────►├─ 查看详情
  ├─ 查看统计分析 ───────►│                         │
  └─ 导出报表             │                         └─''')

# --- Implementation Phases ---
doc.add_heading('10. 分阶段实施计划', level=1)

doc.add_heading('Phase 1 — MVP（2-3周）', level=2)
doc.add_paragraph('目标：题库管理 + 手动判分，跑通核心流程')
add_table(
    ['任务', '优先级', '说明'],
    [
        ['项目脚手架', 'P0', 'React + Express + Prisma + Docker Compose'],
        ['用户认证(JWT)', 'P0', '登录/注册，三角色区分'],
        ['用户管理', 'P0', '管理员 CRUD'],
        ['题目 CRUD', 'P0', '创建、编辑、删除、列表、分类'],
        ['验证规则编辑器', 'P0', '至少支持 check_table_exists, check_field, check_view_exists'],
        ['考试 CRUD', 'P0', '创建考试、选题组卷、发布/开始/结束'],
        ['学生考试流程', 'P0', '查看考试 → 答题说明 → 标记完成 → 提交'],
        ['手动判分界面', 'P0', '教师逐题查看并打分、写评语'],
        ['基础成绩展示', 'P1', '学生查看总分和每题得分'],
    ]
)

doc.add_heading('Phase 2 — 自动判分（2-3周）', level=2)
doc.add_paragraph('目标：集成金山 API，实现自动判分引擎')
add_table(
    ['任务', '优先级', '说明'],
    [
        ['金山API适配器', 'P0', '封装API调用，处理认证、错误重试'],
        ['Rule Engine', 'P0', '规则解释器 + VerifierDispatcher'],
        ['API Verifier', 'P0', '实现全部18种 VerificationAction'],
        ['异步判分队列', 'P0', 'Redis Bull 队列，批量判分'],
        ['人工复核界面', 'P0', 'needs_review标记推送教师'],
        ['Snapshot Verifier', 'P1', '截图对比(Puppeteer + pixelmatch)'],
    ]
)

doc.add_heading('Phase 3 — 考试流程完善（2周）', level=2)
add_table(
    ['任务', '优先级', '说明'],
    [
        ['考试计时', 'P0', '倒计时、到时自动提交'],
        ['WebSocket监控', 'P0', '在线状态、答题进度实时'],
        ['防作弊', 'P1', 'IP限制、切屏检测'],
        ['随机出卷', 'P1', '题目随机排列'],
        ['表格空间分配', 'P1', 'API为每个学生创建独立空间'],
    ]
)

doc.add_heading('Phase 4 — 分析与优化（1-2周）', level=2)
add_table(
    ['任务', '优先级', '说明'],
    [
        ['成绩分析仪表盘', 'P0', '分数分布、正确率、难度分析'],
        ['学生能力画像', 'P1', '知识点雷达图'],
        ['报表导出', 'P1', 'PDF / Excel'],
        ['性能优化', 'P2', '并发判分、缓存策略'],
    ]
)

# --- Key Decisions ---
doc.add_heading('11. 关键设计决策', level=1)
add_table(
    ['决策项', '选择', '理由'],
    [
        ['判分时机', '提交后异步判分', '避免同步阻塞，API调用可能较慢'],
        ['规则存储', 'JSONB存于题目', '灵活扩展，不同题型规则结构差异大'],
        ['人工vs自动', '先自动判，不确定推人工', '平衡效率与准确度'],
        ['题目与考试', '题库独立，考试引用', '题目可复用，组卷灵活'],
        ['工作空间', '每次考试独立空间', '避免干扰，便于清理'],
        ['状态管理', 'Zustand', '非超复杂应用，够用且轻量'],
        ['题型设计', '5种题型', '建表/字段/视图/表单/综合，覆盖主要操作'],
    ]
)

# --- Directory Structure ---
doc.add_heading('12. 项目目录结构', level=1)
add_code_block('''examination-system/
├── docker-compose.yml
├── package.json                     # monorepo root
├── packages/
│   ├── server/
│   │   ├── src/
│   │   │   ├── index.ts             # 入口
│   │   │   ├── app.ts               # Express配置
│   │   │   ├── config/              # 环境配置 + Prisma客户端
│   │   │   ├── middleware/          # auth(JWT), error-handler
│   │   │   ├── routes/              # auth, users, questions, exams, grading, statistics
│   │   │   ├── services/            # 业务逻辑层
│   │   │   ├── engine/              # 验证引擎 ★
│   │   │   │   ├── index.ts         # GradingService
│   │   │   │   ├── rule-engine.ts   # RuleEngine
│   │   │   │   ├── verifiers/       # API/Snapshot/Manual
│   │   │   │   └── adapters/        # kingsoft-adapter.ts
│   │   │   ├── jobs/                # 异步判分队列
│   │   │   └── utils/
│   │   ├── prisma/ (schema + seed)
│   │   └── tests/
│   │
│   └── client/
│       └── src/
│           ├── App.tsx, main.tsx
│           ├── pages/ (auth/, student/, teacher/, admin/)
│           ├── components/ (layout/, common/, rule-builder/)
│           ├── stores/  (auth.ts — Zustand)
│           ├── services/ (api.ts — Axios)
│           └── types/   (index.ts — 全部类型定义)
│
└── docs/ (architecture.md, api-reference.md, rule-engine.md)''')

# --- Verification ---
doc.add_heading('13. 验证方式', level=1)
doc.add_paragraph('开发阶段验证：', style='List Bullet')
for item in [
    '后端单元测试（Vitest）：验证规则引擎每种 action 的判分逻辑',
    'API 集成测试（Supertest）：验证所有 REST 端点',
    '前端组件测试（React Testing Library）：核心组件（RuleBuilder, ExamTimer）',
    'E2E 测试（Playwright）：完整考试流程端到端',
]:
    doc.add_paragraph(item, style='List Bullet 2')

doc.add_paragraph('部署后验证：', style='List Bullet')
for item in [
    '教师创建一道简单题目（建表 + 添字段）→ 发布考试',
    '学生在金山表格中完成操作 → 提交答卷',
    '系统自动判分 → 检查判分结果与预期一致性',
    '教师复核 → 修正误判 → 最终成绩',
]:
    doc.add_paragraph(item, style='List Bullet 2')

# ============================================================
# Save
# ============================================================
output_path = r'C:\Users\AI\ClaudeCode\17_temp\wps_table.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
